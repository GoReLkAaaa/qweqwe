import json
from docx import Document


# user = {
#     "name": "Nikita",
#     "age": 30,
#     "skills": ["Python", "Java"],
#     "languages": {
#         "name": "Russia",
#         "python": "Python",
#         "java": "Java"
#     }
# }

#
# with open('user.json', 'w') as outfile:
#     json.dump(user, outfile, indent=4, ensure_ascii=False)


# with open('user.json', 'r') as f:
#     data = json.load(f)
#     print(data)

#
# load_json_str = json.dumps(user)
# print(load_json_str)
#
#
# text = '{"name" : "Nikita"}'
#
# load_text = json.loads(text)
# print(load_text)



doc = Document()
doc.add_heading("Отчет по проекту", level=1)
doc.add_paragraph("Новый отчет по проекту за 17.11.2025")
para = doc.add_paragraph()
run = para.add_run("Текст рядом")
run.bold = True


table = doc.add_table(rows=2, cols=2)
table.rows[0].cells[0].text = "Имя"
table.rows[0].cells[1].text = "Возраст"
table.rows[1].cells[0].text = "Никита"
table.rows[1].cells[1].text = "23"

doc.add_picture("photo.jpg")


doc.save("report.docx")

for paragraph in doc.paragraphs:
    print(paragraph.text)